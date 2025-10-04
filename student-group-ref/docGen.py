import pandas as pd
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.shared import Inches
from docx.shared import Cm
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH # Specifies paragraph justification type.
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties
import os

# 获取当前文件夹路径
current_dir = os.getcwd()
    
# 初始化一个列表，用于存储文件名中 LS.xlsx 之前的字段
prefix_list = []
    
# 遍历当前文件夹中的所有文件
for filename in os.listdir(current_dir):
    # 检查文件是否以 'LS.xlsx' 结尾
    if filename.endswith('LS.xlsx'):
        # 提取 'LS.xlsx' 之前的部分
        prefix = filename[:-len('LS.xlsx')]
        prefix_list.append(prefix)
        print(f"找到文件：{filename}")
        print(f"提取的字段：{prefix}")
    
# 在当前文件夹内创建名为“报告”的文件夹
report_folder = os.path.join(current_dir, '报告')
if not os.path.exists(report_folder):
    os.makedirs(report_folder)
    print(f"已创建文件夹：{report_folder}")
else:
    print(f"文件夹已存在：{report_folder}")

# 读取Excel文件
class_surfix = '282412244'
df = pd.read_excel(prefix + 'LS.xlsx') # ,nrows=2
print(df.columns)
# 遍历每一行数据
for index, row in df.iterrows():
    # 创建一个新的Word文档
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 向文档中添加内容
    doc.add_heading(str(row['姓名'])+'学习风格报告', 0)

    # 增加一个小节，设置为2栏
    section = doc.add_section(start_type=WD_SECTION_START.CONTINUOUS)
    # section = doc.sections[0]
    section._sectPr.xpath('./w:cols')[0].set(qn('w:num'),'2')
    
    doc.add_paragraph('学号：' + str(row['学号']))

    pic_path = prefix + '_附件/' + '序号' + str(row['序号']) + '_' + str(row['姓名']) + '.jpeg'
##    print(pic_path)
    if os.path.exists(pic_path):
        # 插入图片，并仅设置宽度（高度将自动调整以维持纵横比）       
        doc.add_picture(pic_path, height=Inches(1.8))

    p = doc.add_paragraph('本报告基于 Index of Learning Styles 问卷调查的结果生成，\
可为制定学习策略提供参考。基于 Felder-Silverman 学习风格模型，学习风格可划分为以下四个维度：\
积极/沉思，感官/直觉，视觉/言语，顺序/全局。\
每个学习者在每个维度上的倾向性都有所不同，因此需要制定个性化的学习策略。学习风格的各个维度及其特点如下表：')
    p_format = p.paragraph_format
    p_format.first_line_indent = Inches(0.3)  # 段前空格（0.5英寸）
    p = doc.add_paragraph()

    # 增加一个小节，设置为1栏
    section = doc.add_section(start_type=WD_SECTION_START.CONTINUOUS)
    # section = doc.sections[1]
    section._sectPr.xpath('./w:cols')[0].set(qn('w:num'),'1')
    
    
    
    # 添加一个9x3的表格
    table = doc.add_table(rows=9, cols=3, style='Table Grid')

    # 设置列宽
    for trow in table.rows:
        cell = trow.cells[0]
        cell.width = Cm(2.5)  # 这里设置为3厘米，你可以根据需要调整
        cell = trow.cells[1]
        cell.width = Cm(2)  # 这里设置为3厘米，你可以根据需要调整
        cell = trow.cells[2]
        cell.width = Cm(14.93)  # 这里设置为3厘米，你可以根据需要调整

    # 访问需要合并的单元格，并设置grid_span属性
    cell_00 = table.cell(0, 0)
    cell_01 = table.cell(0, 1)
    cell_00_01 = cell_00.merge(cell_01)
    run = cell_00_01.paragraphs[0].add_run("学习风格类型")
    cell_00_01.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER # 水平居中
##    cell_00_01.paragraphs[0].paragraph_font_name = '黑体'
    run.font.name = '黑体'
    run.font.bold = 1
    run = table.cell(0, 2).paragraphs[0].add_run("特点")
    run.font.name = '黑体'
    run.font.bold = 1
    table.cell(0, 2).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER # 水平居中

    cell_10 = table.cell(1, 0)
    cell_20 = table.cell(2, 0)
    cell_10_20 = cell_10.merge(cell_20)
    cell_10_20.text = "信息加工"
    cell_10_20.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中

    table.cell(1, 1).text = "积极型"
    table.cell(1, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    table.cell(1, 2).text = "通过积极讨论、积极动手和解释给别人听来获取知识；乐于尝试；喜欢协同学习。"
    table.cell(2, 1).text = "沉思型"
    table.cell(2, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    table.cell(2, 2).text = "安静思考；三思而后行；偏爱独自学习或与固定的搭档共同学习；擅长理论。"

    cell_30 = table.cell(3, 0)
    cell_40 = table.cell(4, 0)
    cell_30_40 = cell_30.merge(cell_40)
    cell_30_40.text = "信息感知"
    cell_30_40.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    
    table.cell(3, 1).text = "感官型"
    table.cell(3, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    table.cell(3, 2).text = "擅长记忆事实；不喜欢复杂和意外；对细节很有耐心；更加实际和谨慎；\
喜欢与现实世界有所联系的知识。"
    table.cell(4, 1).text = "直觉型"
    table.cell(4, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    table.cell(4, 2).text = "擅长发现事物之间的联系；喜欢创新，不喜欢重复；擅长掌握新概念及抽象概念；\
对细节较为粗心。"

    cell_50 = table.cell(5, 0)
    cell_60 = table.cell(6, 0)
    cell_50_60 = cell_50.merge(cell_60)
    cell_50_60.text = "信息输入"
    cell_50_60.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    
    table.cell(5, 1).text = "视觉型"
    table.cell(5, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    table.cell(5, 2).text = "喜欢通过可视化的学习资源来获取知识，如视频、图表、概念图等。"
    table.cell(6, 1).text = "言语型"
    table.cell(6, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    table.cell(6, 2).text = "喜欢书面或者口头阐释的学习资源，文本、音频等。"

    cell_70 = table.cell(7, 0)
    cell_80 = table.cell(8, 0)
    cell_70_80 = cell_70.merge(cell_80)
    cell_70_80.text = "内容理解"
    cell_70_80.vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    
    table.cell(7, 1).text = "顺序型"
    table.cell(7, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    table.cell(7, 2).text = "喜欢按照逻辑顺序进行学习；依靠部分信息就可以开展工作。"
    table.cell(8, 1).text = "全局型"
    table.cell(8, 1).vertical_alignment = WD_ALIGN_VERTICAL.CENTER # 垂直居中
    table.cell(8, 2).text = "喜欢从整体角度看待问题，会比较倾向于先掌握知识整体的框架，然后再进行深入学习；\
思维比较活跃和发散。"         
    
    doc.add_heading('1 总览', 1)
    p = doc.add_paragraph('您的学习风格在四个维度下的倾向性为：' )
    p_format = p.paragraph_format
    p_format.first_line_indent = Inches(0.3)  # 段前空格（0.5英寸）
    text = '• 信息加工：'
    if abs(row['积极/沉思']) >= 9:
        text = text + '（强）'
    else:
        if abs(row['积极/沉思']) <= 3:
            text = text + '（弱）'
        else:
            text = text + '（中）'
    if row['积极/沉思'] < 0:
        text = text + '积极'
    else:
        text = text + '沉思'
    doc.add_paragraph(text)

    text = '• 信息感知：'
    if abs(row['感官/直觉']) >= 9:
        text = text + '（强）'
    else:
        if abs(row['感官/直觉']) <= 3:
            text = text + '（弱）'
        else:
            text = text + '（中）'
    if row['感官/直觉'] < 0:
        text = text + '感官'
    else:
        text = text + '直觉'
    doc.add_paragraph(text)

    text = '• 信息输入：'
    if abs(row['视觉/言语']) >= 9:
        text = text + '（强）'
    else:
        if abs(row['视觉/言语']) <= 3:
            text = text + '（弱）'
        else:
            text = text + '（中）'
    if row['视觉/言语'] < 0:
        text = text + '视觉'
    else:
        text = text + '言语'
    doc.add_paragraph(text)

    text = '• 内容理解：'
    if abs(row['顺序/全局']) >= 9:
        text = text + '（强）'
    else:
        if abs(row['顺序/全局']) <= 3:
            text = text + '（弱）'
        else:
            text = text + '（中）'
    if row['顺序/全局'] < 0:
        text = text + '顺序'
    else:
        text = text + '全局'
    doc.add_paragraph(text)

    doc.add_paragraph('（强弱程度仅仅代表倾向性，不代表学习能力强弱）' )

    doc.add_heading('2 学习风格维度', 1)
##    p = doc.add_paragraph('数字标识出了您的学习风格在每个维度下的倾向程度，数字越负越倾向于左侧，\
##越正越倾向于右侧。')
    p = doc.add_paragraph('数字标识出了您的学习风格在每个维度下的倾向程度，在每个维度下总体的倾向取决于两个方向之差。')
    p_format = p.paragraph_format
    p_format.first_line_indent = Inches(0.3)  # 段前空格（0.5英寸）

    # 设置Matplotlib字体
    pfont = FontProperties(fname=r'/System/Library/Fonts/Supplemental/Songti.ttc', size=12)  # 请替换为你系统中的SimSun字体路径

    fig, ax = plt.subplots(figsize=(7,3.5))
##    row_data = [row['积极/沉思'],row['感官/直觉'],row['视觉/言语'],row['顺序/全局']]   
    row_data = [-row['积极'],-row['感官'],-row['视觉'],-row['顺序']]
    bars = ax.barh([4,3,2,1], row_data, color='skyblue')
    row_data = [row['沉思'],row['直觉'],row['言语'],row['全局']]
    bars = ax.barh([4,3,2,1], row_data, color='salmon')

    # 设置y轴的标签为LS列名
    ax.set_yticks([])
    left_yticklabels = ['积极','感官','视觉','顺序']
    right_yticklabels = ['沉思','直觉','言语','全局']


    # 在柱状图的左侧添加标签
    for i, bar in enumerate(bars):
        width = bar.get_width()
        ax.text(-14,  # x位置，偏移量为最大值的5%以避免重叠
                bar.get_y() + bar.get_height() / 2,  # y位置
                left_yticklabels[i],  # 标签文本
                va='center', fontproperties=pfont)  # 垂直对齐方式，字体

    # 在柱状图的右侧添加标签
    for i, bar in enumerate(bars):
        width = bar.get_width()
        ax.text(13,  # x位置，偏移量为最大值的5%以避免重叠
                bar.get_y() + bar.get_height() / 2,  # y位置
                right_yticklabels[i],  # 标签文本
                va='center', fontproperties=pfont)  # 垂直对齐方式，字体

    plt.xlim(-12,12)
    
    plt.xlabel('倾向性', fontproperties=pfont)
##    plt.ylabel('纵轴标签', fontproperties=pfont)
    plt.tight_layout()
    plt.grid(axis='x')

##    plt.title('折线图', fontproperties=pfont)
##    plt.legend(prop=pfont)

    # 保存折线图为图片文件
    plt.savefig('line_chart.png', bbox_inches='tight')
    plt.close()

    # 插入折线图到Word文档
    doc.add_picture('line_chart.png', width=Inches(6.0), height=Inches(3.0))


##    p = doc.add_paragraph('积极   ⟵   ' + str(row['积极/沉思']) + '   ⟶   沉思')
##    paragraph_format = p.paragraph_format
##    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
##    p = doc.add_paragraph('感官   ⟵   ' + str(row['感官/直觉']) + '   ⟶   直觉')
##    paragraph_format = p.paragraph_format
##    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
##    p = doc.add_paragraph('视觉   ⟵   ' + str(row['视觉/言语']) + '   ⟶   言语')
##    paragraph_format = p.paragraph_format
##    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
##    p = doc.add_paragraph('顺序   ⟵   ' + str(row['顺序/全局']) + '   ⟶   全局')
##    paragraph_format = p.paragraph_format
##    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
##    doc.add_paragraph('• 如果您在一个维度的倾向性在 1~3之间，意味着在该维度下的两种特质之间有较好的平衡，\
##因此能够适应不同的学习策略和教学风格。')
##    doc.add_paragraph('• 如果倾向性为 5 或 7，您在该维度下具有中等的倾向性，\
##因此能够更加轻松地在有利于该特质的教学环境中学习。')
##    doc.add_paragraph('• 如果倾向性为 9 或 11，您对该维度下的特质具有很强的倾向性，\
##在不支持这种倾向的环境中学习可能具有较大的困难，因此需要特别关注个性化的策略。')
    doc.add_paragraph('• 如果您在一个维度的倾向性差异在 1~3之间，意味着在该维度下的两种特质之间有较好的平衡，\
因此能够适应不同的学习策略和教学风格。')
    doc.add_paragraph('• 如果倾向性差异为 5 或 7，您在该维度下某个特质具有中等的倾向性，\
因此能够更加轻松地在有利于该特质的教学环境中学习。')
    doc.add_paragraph('• 如果倾向性差异为 9 或 11，您对该维度下的某个特质具有很强的倾向性，\
在不支持这种倾向的环境中学习可能具有较大的困难，因此需要特别关注个性化的策略。')

    doc.add_heading('3 推荐学习策略', 1)
    text = '• '
    if abs(row['积极/沉思']) >= 9:
        text = text + '（强）'
    else:
        if abs(row['积极/沉思']) <= 3:
            text = text + '（弱）'
        else:
            text = text + '（中）'
    if row['积极/沉思'] < 0:
        text = text + '积极学习者：如果您是积极型学习者，在课堂上很少或根本没有时间进行讨论或进行解决问题的活动，\
那么您应该在学习时尝试弥补这些不足。例如在小组中学习，成员轮流向彼此讲解不同的主题。与其他人一起猜测下一次考试会考到什么，\
并弄清楚您将如何作答。如果您找到解决的方法，您将总是更好地掌握接收的信息。'
    else:
        text = text + '沉思学习者：如果您是沉思型学习者，在课堂上很少或根本没有时间来思考新信息，\
那么您应该在学习时设法弥补这一不足。不要只是简单的阅读或记住材料；定期停下来回顾已读内容并考虑可能存在的问题或其应用。\
您可能会发现用自己的话写一些简短的阅读摘要或课堂笔记会很有用。这样做可能会花费额外的时间，但可以使您更有效地掌握材料。'
    doc.add_paragraph(text)

    text = '• '
    if abs(row['感官/直觉']) >= 9:
        text = text + '（强）'
    else:
        if abs(row['感官/直觉']) <= 3:
            text = text + '（弱）'
        else:
            text = text + '（中）'
    if row['感官/直觉'] < 0:
        text = text + '感官学习者：如果感官型学习者能够看到信息与现实世界的联系，则他们最能理解并记住信息。\
如果您所在的班级中大部分教学材料都是抽象和理论性的，那么您可能会遇到学习困难。向您的老师询问概念和步骤的具体示例，\
并了解如何在实践中应用这些概念。如果老师没有提供足够的细节，请尝试在您的课程教科书或其他参考资料中找到一些细节，\
或者与朋友或同学一起集思广益。'
    else:
        text = text + '直觉学习者：许多大学讲课都是针对直觉型学习者的。但是，如果您是一位直觉型学习者，\
并且碰巧参加了主要涉及记忆和死记硬背公式的课程，那么您可能会感到无聊。向您的老师寻求与事实相联系的解释或理论，\
或尝试自己找到联系。您可能还容易在测试中犯粗心大意的错误，因为您对细节不耐烦并且不喜欢重复（例如检查答案）。\
在开始回答之前，请花一些时间阅读整个问题，并确保检查结果。'
    doc.add_paragraph(text)

    text = '• '
    if abs(row['视觉/言语']) >= 9:
        text = text + '（强）'
    else:
        if abs(row['视觉/言语']) <= 3:
            text = text + '（弱）'
        else:
            text = text + '（中）'
    if row['视觉/言语'] < 0:
        text = text + '视觉学习者：如果您是视觉学习者，而课程材料主要以言语形式表达的话，请尝试查找图表，草图，\
示意图，照片，流程图或课程材料的任何其他视觉表示形式。询问您的老师，查阅参考书，并查看课程资料是否有任何视频。\
可以准备概念图：通过列出关键点，将其括在方框或圆圈中并在概念之间用箭头画线以显示连接。用荧光笔对笔记进行颜色编码，\
以相同的颜色标注与一个主题相关的所有内容。'
    else:
        text = text + '言语学习者：用自己的文字写出课程材料的摘要或提纲。小组合作特别有效：通过听取同学的解释，\
您可以对材料有所了解，而在进行解释时，您会学到更多。'
    doc.add_paragraph(text)

    text = '• '
    if abs(row['顺序/全局']) >= 9:
        text = text + '（强）'
    else:
        if abs(row['顺序/全局']) <= 3:
            text = text + '（弱）'
        else:
            text = text + '（中）'
    if row['顺序/全局'] < 0:
        text = text + '顺序学习者：大多数大学课程是按顺序教授的。但是，如果您是顺序学习者，\
并且有一位老师从一个主题跳到另一个主题或跳过一些步骤，那么您可能很难跟随和记忆。要求老师补上跳过的步骤，\
或通过查阅参考资料自行补上。在学习时，请花时间按逻辑顺序为自己概述课程材料。从长远来看，这样做可以节省您的时间。\
您还可以尝试通过将学习的每个新主题与已知知识相关联来增强全局思维能力。您可以做的越多，对主题的理解就可能越深入。'
    else:
        text = text + '全局学习者：如果您是一名全局学习者，认识到自己在掌握细节之前，需要先了解主题的全貌是很有\
帮助的。如果您的老师直接投入新的话题而又不花时间去解释它们与您已经知道的事情之间的关系，那么这可能会给您带来麻烦。\
幸运的是，您可以采取一些步骤来帮助您更快地了解全貌。在开始学习书中章节的第一部分之前，请浏览整个章节以获取概述。\
这样做一开始可能很耗时，但可以避免以后再一遍遍读单个部分。您可能会发现，用大块的时间使自己沉浸在单个主题中会比每晚在\
每个主题花很少时间更有效果。尝试要求老师帮助您理解联系或查阅参考文献，将主题与您已经知道的事情联系起来。最重要的是，不要对自己失去信心。您最终将了解新材料，一旦您了解了它如何与其他主题和学科联系起来，便可以使您以大多数顺序思考者难以想象的方式应用它。'
    doc.add_paragraph(text)
    
##    doc.add_paragraph('• ' + str(row['LSD1']))
##    doc.add_paragraph('• ' + str(row['LSD2']))
##    doc.add_paragraph('• ' + str(row['LSD3']))
##    doc.add_paragraph('• ' + str(row['LSD4']))

    # 保存Word文档
    doc.save(f"报告/序号{row['序号']}_{row['学号']}{row['姓名']}.docx")
